s='君自故乡来，应知故乡事。<br/>来日绮窗前，寒梅著花未？'
a=s.replace('<br/>','\n')
print(s,a)

# with open('h:/test/3.doc','w+') as f:
#     # f.write(a)
#     f.write('title')
#
#     f.write(a)
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

#打开文档
# document = Document('H:/test/wangwei1.docx')
#
# import  docxUtil
# title='寒梅著花'
# document=docxUtil.write_docx(document,title,a,a)
# document.save('H:/test/wangwei3.docx')
# document.styles['Normal'].font.name = '宋体'
# document.styles['Normal'].font.size=Pt(18)
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#标题
# document.add_paragraph(title)
#添加文本
# document.add_paragraph(a)
#
#
#
# document.add_paragraph(title)
# document.add_paragraph(a)
#设置字号
# run = paragraph.add_run(u'设置字号、')
# run.font.size = Pt(24)
#设置中文字体
# run = paragraph.add_run(u'设置中文字体、')
# run.font.name=u'宋体'
#增加分页
# document.add_page_break()
#保存文件
# f_name='h:/test/t10.docx'
# document.save(f_name)

# with open('h:/test/6.pdf', 'w+') as f:
#     # f.write(a)
#     f.write('title'+'\n')
#
#     f.write(a)

total=407
toalPage=int(total/40)+1
print(toalPage)