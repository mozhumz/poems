# from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

def write_docx(document,title,text,sx_str,i):

    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal'].font.size = Pt(18)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 标题
    document.add_paragraph(str(i)+'.'+title)
    # 内容
    document.add_paragraph(text)
    if sx_str:
        document.add_paragraph(sx_str)
    # 增加分页
    document.add_page_break()

    return document


